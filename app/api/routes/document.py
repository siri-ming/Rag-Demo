"""
文档上传和管理路由模块。

提供以下 REST API 端点：
- POST   /api/documents/upload   — 上传并处理文档
- GET    /api/documents           — 获取文档列表
- DELETE /api/documents/{doc_id}  — 删除单个文档
- DELETE /api/documents           — 清空整个知识库（删除所有文档和向量）
- GET    /api/documents/stats     — 获取系统统计信息
"""

import base64
import io
import logging
import os
from pathlib import Path

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse

from app.api.schemas import (
    ClearAllResponse,
    CollectionInfo,
    CollectionListResponse,
    DeleteResponse,
    DocumentListItem,
    DocumentListResponse,
    StatsResponse,
    UploadResponse,
    ChunkInfo,
    ChunkListResponse,
)
from app.config import settings
from app.parsers import get_supported_extensions
from app.services.document_service import get_document_service

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/documents", tags=["documents"])


@router.post("/upload", response_model=UploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    collection: str = Form(default="default"),
    chunk_strategy: str = Form(default=None),
    chunk_size: int = Form(default=None),
    chunk_overlap: int = Form(default=None),
    sliding_window: int = Form(default=0),
):
    """
    上传并处理文档。

    文件将被保存到上传目录，然后自动经过解析→分块→向量化→存储的完整流水线。

    参数:
        file: 上传的文件（支持 PDF、Word、Excel、图片等格式）。
        collection: 知识库名称，默认为"default"。
        chunk_strategy: 可选的分块策略（fixed / recursive / parent_child）。
        chunk_size: 可选的分块大小。
        chunk_overlap: 可选的分块重叠。
        sliding_window: 滑动窗口大小，0表示不使用。

    返回:
        文档处理结果，包含 doc_id、分块数等元数据。
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="未提供文件名")

    # 校验文件类型
    ext = Path(file.filename).suffix.lower()
    supported = get_supported_extensions()
    if ext not in supported:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型: '{ext}'。支持的格式: {', '.join(supported)}",
        )

    # 读取文件内容并校验大小
    content = await file.read()
    if len(content) > settings.max_upload_size:
        raise HTTPException(
            status_code=400,
            detail=f"文件过大。最大允许: {settings.max_upload_size // (1024*1024)}MB",
        )

    # 保存文件到上传目录
    os.makedirs(settings.upload_dir, exist_ok=True)
    file_path = os.path.join(settings.upload_dir, file.filename)

    # 处理文件名冲突：自动追加序号
    counter = 1
    base_name = Path(file.filename).stem
    while os.path.exists(file_path):
        file_path = os.path.join(settings.upload_dir, f"{base_name}_{counter}{ext}")
        counter += 1

    with open(file_path, "wb") as f:
        f.write(content)

    # 调用文档服务处理文件
    try:
        service = get_document_service()
        result = service.process_document(
            file_path=file_path,
            collection=collection,
            chunk_strategy=chunk_strategy,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            sliding_window=sliding_window if sliding_window > 0 else None,
        )
        return UploadResponse(**result)
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        logger.exception("文档处理失败: %s", file.filename)
        raise HTTPException(
            status_code=500,
            detail=f"文档处理失败: {str(e)}",
        )


@router.get("", response_model=DocumentListResponse)
async def list_documents():
    """获取所有已上传文档的列表。"""
    service = get_document_service()
    docs = service.list_documents()
    items = [DocumentListItem(**doc) for doc in docs]
    return DocumentListResponse(documents=items, total=len(items))


@router.delete("/{doc_id}", response_model=DeleteResponse)
async def delete_document(doc_id: str):
    """
    删除单个文档及其关联的向量数据。

    删除 documents 表记录时，document_chunks 表中的分块会通过外键级联自动删除。
    """
    service = get_document_service()
    success = service.delete_document(doc_id)

    if not success:
        raise HTTPException(status_code=404, detail=f"文档不存在: {doc_id}")

    return DeleteResponse(
        success=True,
        doc_id=doc_id,
        message="文档删除成功",
    )


@router.delete("", response_model=ClearAllResponse)
async def clear_all_documents():
    """
    清空整个知识库。

    删除所有文档记录和所有向量分块，此操作不可撤销。
    上传的原始文件不会被删除（仍保留在 uploads 目录中）。
    """
    service = get_document_service()
    result = service.clear_all()

    return ClearAllResponse(
        success=True,
        message=f"知识库已清空：删除了 {result['deleted_documents']} 篇文档，{result['deleted_chunks']} 个分块",
        deleted_documents=result["deleted_documents"],
        deleted_chunks=result["deleted_chunks"],
    )


@router.get("/stats", response_model=StatsResponse)
async def get_stats():
    """获取系统统计信息（文档总数、支持格式、分块策略等）。"""
    service = get_document_service()
    return StatsResponse(**service.get_stats())


@router.get("/collections", response_model=CollectionListResponse)
async def list_collections():
    """获取所有知识库列表及其统计信息。"""
    service = get_document_service()
    collections = service.list_collections()
    items = [CollectionInfo(**c) for c in collections]
    return CollectionListResponse(collections=items, total=len(items))


@router.get("/{doc_id}/chunks", response_model=ChunkListResponse)
async def get_document_chunks(doc_id: str):
    """
    获取指定文档的所有分块详情。
    
    用于前端预览文档的切分效果，验证分块策略是否达到预期。
    """
    service = get_document_service()
    # 先检查文档是否存在
    doc = service.get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail=f"文档不存在: {doc_id}")
    
    chunks = service.get_document_chunks(doc_id)
    items = [ChunkInfo(**c) for c in chunks]
    return ChunkListResponse(doc_id=doc_id, chunks=items, total=len(items))


@router.get("/{doc_id}/preview")
async def preview_source_file(doc_id: str):
    """
    预览源文件。

    返回上传的原始文件，支持浏览器内联显示（PDF/图片/文本等）。
    """
    service = get_document_service()
    doc = service.get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail=f"文档不存在: {doc_id}")

    file_name = doc["file_name"]
    file_path = os.path.join(settings.upload_dir, file_name)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="源文件不存在")

    # 根据文件类型设置 media_type
    ext = Path(file_name).suffix.lower()
    media_types = {
        ".pdf": "application/pdf",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".txt": "text/plain; charset=utf-8",
        ".md": "text/plain; charset=utf-8",
        ".csv": "text/plain; charset=utf-8",
        ".json": "application/json; charset=utf-8",
        ".html": "text/html; charset=utf-8",
        ".xml": "application/xml; charset=utf-8",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }
    media_type = media_types.get(ext, "application/octet-stream")

    return FileResponse(
        path=file_path,
        media_type=media_type,
        filename=file_name,
        content_disposition_type="inline",
    )


# ---------- DOCX / XLSX → HTML 预览 ----------

_DOCX_PREVIEW_CSS = """
body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
       line-height: 1.8; color: #1e293b; max-width: 860px; margin: 0 auto; padding: 32px 24px;
       background: #fff; }
h1,h2,h3,h4 { color: #0f172a; margin-top: 1.4em; }
p { margin: 0.6em 0; }
table { border-collapse: collapse; width: 100%; margin: 1em 0; }
th, td { border: 1px solid #e2e8f0; padding: 8px 12px; text-align: left; font-size: 14px; }
th { background: #f1f5f9; font-weight: 600; }
tr:nth-child(even) { background: #f8fafc; }
img { max-width: 100%; height: auto; border-radius: 6px; margin: 8px 0; }
.heading1 { font-size: 1.8em; font-weight: 700; }
.heading2 { font-size: 1.5em; font-weight: 700; }
.heading3 { font-size: 1.25em; font-weight: 600; }
"""


def _docx_to_html(file_path: str) -> str:
    """将 DOCX 文件转换为 HTML 字符串。"""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    parts: list[str] = []

    # 收集图片 blob → base64 data URI 映射
    image_map: dict[str, str] = {}
    for rel in doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            blob = rel.target_part.blob
            ct = rel.target_part.content_type or "image/png"
            b64 = base64.b64encode(blob).decode()
            image_map[rel.rId] = f"data:{ct};base64,{b64}"
        except Exception:
            pass

    # 构建 body 元素顺序映射（段落和表格按文档顺序）
    from docx.oxml.ns import qn
    body = doc.element.body

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # 段落
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue

            text = para.text.strip()
            style_name = (para.style.name or "").lower() if para.style else ""

            # 检查是否有内联图片 (run 中的 drawing)
            has_image = False
            img_html = ""
            for run in para.runs:
                drawings = run._element.findall(qn("w:drawing"))
                for drawing in drawings:
                    blips = drawing.findall(".//" + qn("a:blip"))
                    for blip in blips:
                        embed = blip.get(qn("r:embed"))
                        if embed and embed in image_map:
                            img_html += f'<img src="{image_map[embed]}" />'
                            has_image = True

            if not text and not has_image:
                continue

            # 根据样式判断标题级别
            if "heading 1" in style_name or style_name == "title":
                parts.append(f'<h1 class="heading1">{_esc(text)}</h1>')
            elif "heading 2" in style_name:
                parts.append(f'<h2 class="heading2">{_esc(text)}</h2>')
            elif "heading 3" in style_name:
                parts.append(f'<h3 class="heading3">{_esc(text)}</h3>')
            elif "heading" in style_name:
                parts.append(f'<h3>{_esc(text)}</h3>')
            elif "list" in style_name:
                parts.append(f'<li>{_esc(text)}</li>')
            else:
                # 处理加粗/斜体等 run 级别格式
                rich_html = ""
                for run in para.runs:
                    t = _esc(run.text) if run.text else ""
                    if not t:
                        continue
                    if run.bold and run.italic:
                        rich_html += f"<b><i>{t}</i></b>"
                    elif run.bold:
                        rich_html += f"<b>{t}</b>"
                    elif run.italic:
                        rich_html += f"<i>{t}</i>"
                    else:
                        rich_html += t
                if rich_html:
                    parts.append(f"<p>{rich_html}</p>")
                elif text:
                    parts.append(f"<p>{_esc(text)}</p>")

            if img_html:
                parts.append(img_html)

        elif tag == "tbl":
            # 表格
            tbl = None
            for t in doc.tables:
                if t._element is child:
                    tbl = t
                    break
            if tbl is None:
                continue
            parts.append("<table>")
            for i, row in enumerate(tbl.rows):
                parts.append("<tr>")
                cell_tag = "th" if i == 0 else "td"
                for cell in row.cells:
                    parts.append(f"<{cell_tag}>{_esc(cell.text)}</{cell_tag}>")
                parts.append("</tr>")
            parts.append("</table>")

    body_html = "\n".join(parts)
    return (
        f"<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"<style>{_DOCX_PREVIEW_CSS}</style></head>"
        f"<body>{body_html}</body></html>"
    )


def _xlsx_to_html(file_path: str) -> str:
    """将 XLSX 文件转换为 HTML 字符串。"""
    import openpyxl

    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    sheets_html: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        table_parts = [f"<h2>{_esc(sheet_name)}</h2><table>"]
        for i, row in enumerate(rows):
            table_parts.append("<tr>")
            cell_tag = "th" if i == 0 else "td"
            for cell in row:
                val = _esc(str(cell)) if cell is not None else ""
                table_parts.append(f"<{cell_tag}>{val}</{cell_tag}>")
            table_parts.append("</tr>")
        table_parts.append("</table>")
        sheets_html.append("\n".join(table_parts))

    wb.close()
    body = "\n".join(sheets_html) or "<p>空工作簿</p>"
    return (
        f"<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"<style>{_DOCX_PREVIEW_CSS}</style></head>"
        f"<body>{body}</body></html>"
    )


def _esc(text: str) -> str:
    """HTML 转义。"""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


@router.get("/{doc_id}/preview-html")
async def preview_as_html(doc_id: str):
    """
    将 DOCX / XLSX 等无法浏览器原生渲染的文件转为 HTML 预览。
    """
    service = get_document_service()
    doc = service.get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail=f"文档不存在: {doc_id}")

    file_name = doc["file_name"]
    file_path = os.path.join(settings.upload_dir, file_name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="源文件不存在")

    ext = Path(file_name).suffix.lower()

    if ext == ".docx":
        html = _docx_to_html(file_path)
    elif ext in (".xlsx", ".xls"):
        html = _xlsx_to_html(file_path)
    else:
        raise HTTPException(status_code=400, detail=f"不支持 HTML 转换的格式: {ext}")

    return HTMLResponse(content=html)
