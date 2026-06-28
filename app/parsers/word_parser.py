"""Word 文档解析器，基于 python-docx。

支持提取段落、表格以及嵌入图片（通过 OCR 识别图片中的文字）。
"""

import io
import logging
from pathlib import Path

from docx import Document as DocxDocument
from PIL import Image

from .base import BaseParser, Document

logger = logging.getLogger(__name__)

# 图片 OCR 的最小尺寸（像素），过滤装饰性小图标
_MIN_IMAGE_SIZE = 50


class WordParser(BaseParser):
    """
    Microsoft Word (.docx) 文件解析器。

    提取段落文本、表格内容，以及嵌入图片（通过 OCR 识别文字）。
    """

    SUPPORTED_EXTENSIONS = [".docx", ".doc"]

    def __init__(self):
        self._ocr_engine = None

    def _get_ocr_engine(self):
        """懒加载 OCR 引擎。"""
        if self._ocr_engine is None:
            from rapidocr_onnxruntime import RapidOCR
            self._ocr_engine = RapidOCR()
        return self._ocr_engine

    def _extract_and_ocr_images(self, docx_doc: DocxDocument) -> str:
        """
        从 Word 文档中提取嵌入图片并逐一 OCR 识别。

        遍历文档的所有关联关系，找到图片类型并送 OCR。
        过滤太小的图片（装饰性图标等）。

        返回:
            所有图片 OCR 识别出的文本内容。
        """
        engine = self._get_ocr_engine()
        ocr_texts = []

        for rel in docx_doc.part.rels.values():
            if "image" not in rel.reltype:
                continue
            try:
                image_data = rel.target_part.blob
                img = Image.open(io.BytesIO(image_data))

                # 过滤太小的图片
                if img.width < _MIN_IMAGE_SIZE or img.height < _MIN_IMAGE_SIZE:
                    continue

                # 确保 RGB 或灰度模式
                if img.mode not in ("RGB", "L"):
                    img = img.convert("RGB")

                result, _ = engine(img)
                if result:
                    page_text = "\n".join(
                        line[1] for line in result if line[1]
                    )
                    if page_text.strip():
                        ocr_texts.append(page_text.strip())
                        logger.debug(
                            "Word 图片 OCR: %dx%d -> %d 字符",
                            img.width, img.height, len(page_text),
                        )
            except Exception as e:
                logger.debug("Word 图片 OCR 失败: %s", e)
                continue

        return "\n\n".join(ocr_texts)

    def _extract_tables(self, doc: DocxDocument) -> list[str]:
        """提取文档中所有表格的文本内容。"""
        table_texts = []
        for table in doc.tables:
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            if rows_text:
                table_texts.append("\n".join(rows_text))
        return table_texts

    def parse(self, file_path: str | Path) -> list[Document]:
        """
        解析 Word 文档，提取段落、表格和图片 OCR 内容。
        """
        file_path = Path(file_path)
        documents = []

        docx_doc = DocxDocument(str(file_path))
        file_name = file_path.name

        # ---- 提取段落 ----
        paragraph_text = []
        for para in docx_doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraph_text.append(text)

        if paragraph_text:
            documents.append(
                Document(
                    text="\n\n".join(paragraph_text),
                    metadata={
                        "source": file_name,
                        "file_type": "word",
                        "content_type": "paragraphs",
                    },
                )
            )

        # ---- 提取表格 ----
        table_texts = self._extract_tables(docx_doc)
        for i, table_text in enumerate(table_texts):
            documents.append(
                Document(
                    text=table_text,
                    metadata={
                        "source": file_name,
                        "file_type": "word",
                        "content_type": "table",
                        "table_index": i + 1,
                    },
                )
            )

        # ---- 提取嵌入图片并 OCR ----
        try:
            img_text = self._extract_and_ocr_images(docx_doc)
            if img_text:
                documents.append(
                    Document(
                        text=img_text,
                        metadata={
                            "source": file_name,
                            "file_type": "word",
                            "content_type": "image_ocr",
                        },
                    )
                )
                logger.info(
                    "Word 文档图片 OCR: %s, 识别 %d 字符",
                    file_name, len(img_text),
                )
        except Exception as e:
            logger.warning("Word 图片 OCR 整体失败: %s", e)

        return documents
